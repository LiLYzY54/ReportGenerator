
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from latex2mathml.converter import convert as latex2mathml
from mathml2omml import convert as mathml2omml

def append_text_with_math(paragraph, text):
    """Parses text for $...$ and adds it to paragraph as text runs or OMML formulas. Also strips [cite: ...]."""
    if not text:
        return
    # Strip [cite: ...] tags
    text = re.sub(r'\[cite:.*?\]', '', text)
    # Regex to find $...$ (non-greedy)
    parts = re.split(r'(\$.*?\$)', text)
    for part in parts:
        if part.startswith('$') and part.endswith('$'):
            latex = part[1:-1]
            try:
                # 1. LaTeX -> MathML
                mathml = latex2mathml(latex)
                # 2. MathML -> OMML
                omml = mathml2omml(mathml)
                # 3. Insert OMML into Word XML
                # Fix namespace error: Add xmlns:m to the root oMath element
                if 'xmlns:m' not in omml:
                    omml = omml.replace('<m:oMath', '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"')
                paragraph._element.append(parse_xml(omml))
            except Exception as e:
                print(f"Math conversion failed for {latex}: {e}")
                paragraph.add_run(part)
        else:
            if part:
                paragraph.add_run(part)

def set_cell_background(cell, fill):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"})
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for side in ["top", "left", "bottom", "right"]:
        if side in kwargs:
            edge = OxmlElement(f'w:{side}')
            for key, val in kwargs[side].items():
                edge.set(qn(f'w:{key}'), str(val))
            tcBorders.append(edge)

def convert_html_to_docx(html_path, docx_path):
    if not os.path.exists(html_path):
        print(f"Error: {html_path} not found")
        return

    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    # Necessary for Microsoft YaHei to show up correctly in Word
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 1. Header
    header = soup.find(class_=['header', 'lr-header'])
    if header:
        h1 = header.find(['h1', 'div'], class_=['lr-brand-name']) or header.find('h1')
        if h1:
            p_h1 = doc.add_paragraph()
            p_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_h1.add_run(h1.get_text().strip())
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
        info = header.find(['p', 'div'], class_=['lr-date']) or header.find('p')
        if info:
            p_info = doc.add_paragraph()
            p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_info.add_run(info.get_text().strip())
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            
    # 2. Process Cards
    cards = soup.find_all(class_=['card', 'lr-card', 'lr-page'])
    for card in cards:
        # Section Title
        title_tag = card.find(class_=['section-title', 'lr-section-label'])
        if title_tag:
            p_title = doc.add_paragraph()
            run = p_title.add_run(title_tag.get_text().strip())
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        # Diagnosis Box / Evaluation Box
        diag = card.find(class_=['diagnosis-box', 'lr-eval-box'])
        if diag:
            table = doc.add_table(rows=1, cols=1)
            table.width = Inches(6.5)
            cell = table.cell(0, 0)
            set_cell_background(cell, "FEF2F2")
            set_cell_border(cell, left={"sz": 24, "val": "single", "color": "EA580C"})
            
            h3 = diag.find(['h3', 'div'], class_=['lr-eval-text']) or diag.find('h3')
            if h3:
                p_h3 = cell.paragraphs[0]
                run = p_h3.add_run(h3.get_text().strip())
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)
                run.font.size = Pt(11)
            
            p_content = diag.find('p')
            if p_content:
                p_p = cell.add_paragraph()
                append_text_with_math(p_p, p_content.get_text().strip())

        # Grid / Dimension Cards / Stat Cards
        grid = card.find(class_=['grid', 'lr-stat-row', 'lr-card-grid'])
        if grid:
            dim_cards = grid.find_all(class_=['dimension-card', 'lr-stat-card', 'lr-card'])
            if dim_cards:
                grid_table = doc.add_table(rows=1, cols=len(dim_cards))
                grid_table.autofit = True
                for i, dim in enumerate(dim_cards):
                    cell = grid_table.cell(0, i)
                    set_cell_background(cell, "F8FAFC")
                    set_cell_border(cell, 
                                   top={"sz": 4, "val": "single", "color": "E2E8F0"},
                                   bottom={"sz": 4, "val": "single", "color": "E2E8F0"},
                                   left={"sz": 4, "val": "single", "color": "E2E8F0"},
                                   right={"sz": 4, "val": "single", "color": "E2E8F0"})
                    
                    h4 = dim.find(['h4', 'div'], class_=['lr-stat-label']) or dim.find('h4')
                    if h4:
                        p_h4 = cell.paragraphs[0]
                        run = p_h4.add_run(h4.get_text().strip())
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                    
                    p_dim = dim.find(['p', 'div'], class_=['lr-stat-val']) or dim.find('p')
                    if p_dim:
                        p_p = cell.add_paragraph()
                        append_text_with_math(p_p, p_dim.get_text().strip())
                        p_p.paragraph_format.space_before = Pt(6)

        # Standard Data Table
        html_table = card.find('table', class_=['lr-table']) or card.find('table')
        if html_table:
            rows = html_table.find_all('tr')
            if rows:
                header_cells = rows[0].find_all(['th', 'td'])
                doc_table = doc.add_table(rows=len(rows), cols=len(header_cells))
                doc_table.style = 'Table Grid'
                
                for r_idx, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for c_idx, html_cell in enumerate(cells):
                        if c_idx >= len(header_cells): continue
                        doc_cell = doc_table.cell(r_idx, c_idx)
                        
                        # Handle bold text in table (like 第4(6)题)
                        content = html_cell.get_text(separator="\n").strip()
                        doc_cell.text = ""
                        p = doc_cell.paragraphs[0]
                        
                        # Check for bold parts or just take all
                        if not (html_cell.find('b') or html_cell.find('strong') or html_cell.name == 'th'):
                            append_text_with_math(p, content)
                        
                        if html_cell.name == 'th':
                            set_cell_background(doc_cell, "334155")
                            run = p.add_run(content)
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True
                        elif html_cell.find('b') or html_cell.find('strong'):
                            # Simple approach: if any part is bold, we parse children
                            p.clear()
                            for child in html_cell.children:
                                if child.name in ['b', 'strong']:
                                    run = p.add_run(child.get_text())
                                    run.font.bold = True
                                elif isinstance(child, str):
                                    append_text_with_math(p, child)
                                elif child.name == 'br':
                                    p.add_run('\n')
                                else:
                                    append_text_with_math(p, child.get_text())

        # Advice Sections
        advices = card.find_all(class_='advice-section')
        for advice in advices:
            h3 = advice.find('h3')
            if h3:
                p_h3 = doc.add_paragraph()
                run = p_h3.add_run(h3.get_text().strip())
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            
            p_advice = advice.find('p')
            if p_advice:
                p_p = doc.add_paragraph()
                # Handle strong tags inside p
                for child in p_advice.children:
                    if child.name in ['strong', 'b']:
                        run = p_p.add_run(child.get_text())
                        run.font.bold = True
                    elif isinstance(child, str):
                        append_text_with_math(p_p, child)
                    else:
                        append_text_with_math(p_p, child.get_text())

        # Follow-up Learning Planning (Another table)
        # This is already handled by the generic table handler above

    # Add space between cards
    for paragraph in doc.paragraphs:
        if paragraph.text == "":
            paragraph.paragraph_format.space_after = Pt(12)

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

if __name__ == "__main__":
    import argparse
    import os
    
    parser = argparse.ArgumentParser(description="Convert HTML to DOCX")
    parser.add_argument("--html", help="Path to the input HTML file")
    parser.add_argument("--output", help="Path to the output DOCX file")
    
    args = parser.parse_args()
    
    if args.html and args.output:
        convert_html_to_docx(args.html, args.output)
    else:
        # Default behavior for hardcoded files if no args provided
        home_dir = os.path.expanduser("~")
        downloads_dir = os.path.join(home_dir, "Downloads")
        
        files_to_convert = [
            ('lzx-E.html', os.path.join(downloads_dir, '李政兴 英语 近期题目报告.docx')),
        ]
        
        for html_file, docx_file in files_to_convert:
            if os.path.exists(html_file):
                print(f"Converting {html_file}...")
                convert_html_to_docx(html_file, docx_file)
            else:
                print(f"Skipping {html_file} (not found)")
